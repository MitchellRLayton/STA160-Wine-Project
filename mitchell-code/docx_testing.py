# Script for testing docx files that Duncan provided from PDF Pen Pro Results

import pandas as pd
import numpy as np

import os
from PIL import Image, ImageEnhance, ImageFilter, features
import cv2

import docx
import pytesseract
pytesseract.pytesseract.tesseract_cmd = 'C:\\Users\\layto\\AppData\\Local\\Tesseract-OCR\\tesseract'

docxPath = 'C:\\Users\\layto\\SampleWineCatalogs\\Docx_Files'

class wordDocx():
    
    def __init__(self, path):
        self.path = path
    
    def fileList(path):
        '''
        Generates full file paths to Docx files given relative directory input
        '''
        import os
        fullPathFiles = []
        file = os.listdir(path)
        for doc in range(len(file)):
            fullPathFiles.append(path + '\\' + file[doc])
        return(fullPathFiles)
    
    def getAllParagraphs(files):
        import docx
        if isinstance(files,list) is True:
            docPara = []
            for file in files:
                iterDoc = docx.Document(file)
                docText = []
                for par in range(len(iterDoc.paragraphs)):
                    text = iterDoc.paragraphs[par].text
                    makeLines = text.split(' ')
                    if list(textParse.whichIsMore(makeLines).items())[0][0] == '\n':
                        newLines = text.split('\n')
                        docText.append(newLines)
                    elif list(textParse.whichIsMore(makeLines).items())[0][0] == '\t':
                        newLines = text.split('\t')
                        docText.append(newLines)
                listOfText = [y for y in [x[0] if len(x) == 1 else ' '.join(x) for x in docText] if y != '']
                docPara.append(listOfText)
            return(docPara)

        else:
            raise AttributeError('Needs list of full file paths to the .docx files')
        


class textParse():

    def __init__(self, string):
        self.string = string
    
    def findParenth(userInput):
        '''
        Takes a string or a list of strings and finds any parenthesis
        Returns ???
        '''
        if isinstance(userInput,str) is True:
            if string.find('(') != -1 or string.find(')') != -1:
                return(string)
            else:
                return(0)
        elif isinstance(userInput,list) is True:
            found = []
            for line in userInput:
                if line.find('(') != -1 or line.find(')') != -1:
                    found.append(line)
                else:
                    continue
            return(found)
        else:
            raise AttributeError('Only input strings and lists of strings')
        
    def tabInString(string):
        '''
        String input to find tabs: '\t'
        Returns True/False
        '''
        if string.find('\t') != -1:
            return(True)
        else:
            return(False)

    def newLineInString(string):
        '''
        String input to find tabs: '\n'
        Returns True/False
        '''
        if string.find('\n') != -1:
            return(True)
        else:
            return(False)
        
    def whichIsMore(self):
        '''
        Input a list of string (say if you split the text on spaces)
        Returns tuple of counts of '\n' and '\t' respectively
        '''
        if self and isinstance(self,list) is True:
            counterN = 0
            counterT = 0
            for line in range(len(self)):
                if self[line].find('\n') != -1:
                    counterN += 1
                elif self[line].find('\t') != -1:
                    counterT += 1
                elif self[line].find('\n') != -1 and self[line].find('\t') != -1:
                    counterN += 1
                    counterT += 1
            result = (counterN,counterT)
            if result[0] > result[1]:
                return({'\n':result[0],'by':result[0]-result[1]})
            else:
                return({'\t':result[1],'by':result[1]-result[0]})
        else:
            raise AttributeError('Input a list of strings')


class identifyFormat():
    
    def __init__(self, pargraph):
        self.paragraph = paragraph
        
    def years(mini,maxi):
        '''
        Generates a set of guess year ranges for alcohols
        '''
        import numpy as np
        years = set([n for n in np.arange(mini,maxi,1)])
        return(years)

    def removeSpecial(self):
        
        special = set([',','"','|','\\',':',';','`','~','$','#'])
        split = self.split(' ')
        split = [x for x in split if len(x) > 0]
        fixed = []
        
        for i in range(len(split)):
            test = list(split[i])
            if len(set(test).intersection(special)) == 0 and '.' not in test:
                fixed.append(split[i])
            elif test[-1] == '.' and test[0] != '.':
                test = test[:-1]
                split[i] = ''.join(test)
                fixed.append(split[i])
            elif '.' in test:
                if len([x for x in test if x == '.']) == 1:
                    try:
                        if isinstance(int(''.join([x for x in test if x != '.'])),int):
                            split[i] = ''.join(test)
                            fixed.append(split[i])
                    except ValueError:
                        continue
                elif len([y for y in [x.find('.') for x in test] if y == 0]) != 1:
                    length = len([y for y in [x.find('.') for x in test] if y == 0])-1
                    test = test[length:]
                    split[i] = ''.join(test)
                    fixed.append(split[i])
            else:
                while len(set(test).intersection(special)) > 0:
                    ind = [test.index(x) for x in set(test).intersection(special)]
                    for index in ind:
                        test.pop(index)
                    split[i] = ''.join(test)
                    if len(set(split[i]).intersection(special)) == 0:
                        fixed.append(split[i])
                        break
                    else:
                        continue
                
        fixed = ' '.join([''.join(x) for x in fixed])
        return(fixed)
    
    
    def fixBottlePrices(self):
        '''
        Removes special, and fixes bottle prices
        '''
        fixed = identifyFormat.removeSpecial(self)
        split = fixed.split(' ')
        split = [x for x in split if len(x) > 0]
        setRange = np.arange(1,5000,1)
        
        for i in range(len(split)):
            if '.' in list(split[i]) and i >= 2:
                tryInt = split[i-1]
                intList = list(tryInt)
                try:
                    if isinstance(int(tryInt),int):
                        getLength = len(intList)
                        if getLength == 1:
                            split[i-1] = intList[0]+'.'+'00'
                        elif getLength == 2:
                            split[i-1] = intList[0]+'.'+intList[1]+'0'
                        elif getLength == 3:
                            split[i-1] = intList[0]+'.'+intList[1]+intList[2]
                        elif getLength == 4:
                            split[i-1] = intList[0]+intList[1]+'.'+intList[2]+intList[3]
                        else:
                            split[i-1] = split[i-1]
                except ValueError:
                    continue
                    
        for i in range(len(split)+1):
            try:
                single = split[i]
                testNO = split[i-1]
                try:
                    if int(single) in np.arange(1,10,1):
                        if testNO.lower() != 'no':
                            ind = split.index(single)
                            split.pop(int(ind))
                except ValueError:
                    continue
            except IndexError:
                break

        return(' '.join(split))
    
                                
    def findYears(self,mini,maxi):
        '''
        finds year(s) in fixed word docx paragraphs associated with respective alcohol
        '''
        yearSet = identifyFormat.years(mini,maxi)
        fixed = identifyFormat.removeSpecial(self)
        split = fixed.split(' ')
        y = []
        for chars in range(len(split)):
            try:
                year = int(split[chars])
                if year and year in yearSet:
                    y.append(year)
            except ValueError:
                continue
                
        if len(y) > 0:
            return(y)
        else:
            return(0)

    def wordBottle(self):
        '''
        From string of paragraph as input, finds instance of word 'bottle'
        '''
        try:
            test = [x.lower() for x in self.split(' ')]
            find = test.index('bottle')
            if find:
                return(True)
        except ValueError:
            return(False)
        
    def wordCase(self):
        '''
        From string of paragraph as input, finds instance of word 'case'
        '''
        try:
            test = [x.lower() for x in self.split(' ')]
            find = test.index('case')
            if find:
                return(True)
        except ValueError:
            return(False)
        
    def idNum(self):
        '''
        Generates a set of bin ranges for alcohols
        '''
        fixed = self.split(' ')
        import numpy as np
        set1 = set([n for n in np.arange(1,1880,1)])
        set2 = set([n for n in np.arange(2000,15000,1)])
        ID = set(set1.union(set2))
        potentialIDs = []

        for i in fixed:
            try:
                if int(i) in ID:
                    ind = fixed.index(i)
                    potentialIDs.append(ind)
            except ValueError:
                continue
            
        return(potentialIDs)



def finalForm(para):
    # 1) In the case of [no bottle/bin numbers & no year] format accordingly
        # 1.1 [no bottle/bin numbers & with years]
    
    
    
    #2) In the case of [bottle/bin numbers & no year] format accordingly
        # 2.2 [bottle/bin numbers & with years]
    split = para.split(' ')
    print(split)


def mainFunction(allText):
    counter = 0
    for i in range(len(allText)):
        paragraph = allText[i]
        typ = type(paragraph)
        if paragraph != '' and paragraph != '\n' and paragraph != '\t' and paragraph != None:
            if identifyFormat.wordBottle(paragraph) is True or identifyFormat.wordCase(paragraph) is True:
                
                print(identifyFormat.removeSpecial(paragraph))

                # RUN FUNCTIONS ON THIS
                counter += 1
                print(counter)
    
if __name__ == "__main__":
    allText = wordDocx.getAllParagraphs(wordDocx.fileList(docxPath))
    with open('docxParagraphs.txt', 'w') as F:
        F.writelines([str(x[0]) if len(x)==1 else ' '.join(x) + '\n' for x in allText])  

    with open('docxParagraphs.txt','r') as file:
        lines = []
        for line in file:
            lines.append(line)

    for i in range(len(lines)):
        if lines[i].split(' ')[0] == 'Rivero':
            test123 = lines[i]

    print(test123)
    new_test = identifyFormat.fixBottlePrices(test123)
    print('\n',new_test.split('\n')[0])























